"""Create a sample DOCX fixture for testing."""

from pathlib import Path

import docx


def create_sample_docx():
    """Create a DOCX with headings, paragraphs, a table, and an image placeholder."""
    doc = docx.Document()

    doc.core_properties.title = "Test Document"
    doc.core_properties.author = "Test Author"
    doc.core_properties.subject = "Testing"

    doc.add_heading("Introduction", level=1)
    doc.add_paragraph("This is the first paragraph of the test document.")
    doc.add_paragraph("This document contains multiple sections for testing.")

    doc.add_heading("Data Section", level=2)
    doc.add_paragraph("Below is a table with sample invoice data.")

    # Add a table
    table = doc.add_table(rows=4, cols=3)
    table.style = "Table Grid"
    headers = ["Item", "Quantity", "Price"]
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header

    data = [
        ["Widget A", "10", "25.00"],
        ["Widget B", "5", "50.00"],
        ["Total", "", "500.00"],
    ]
    for row_idx, row_data in enumerate(data):
        for col_idx, val in enumerate(row_data):
            table.rows[row_idx + 1].cells[col_idx].text = val

    doc.add_heading("Conclusion", level=2)
    doc.add_paragraph("This concludes the test document.")

    output_dir = Path(__file__).parent / "docx"
    output_dir.mkdir(exist_ok=True)
    output_path = output_dir / "sample.docx"
    doc.save(str(output_path))
    print(f"Created {output_path}")


if __name__ == "__main__":
    create_sample_docx()
