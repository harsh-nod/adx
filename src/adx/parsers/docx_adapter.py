"""DOCX parser adapter using python-docx."""

from __future__ import annotations

import logging
from pathlib import Path
from typing import Any

from adx.models.document import FileType
from adx.parsers.base import (
    ParserAdapter,
    ParserCapabilities,
    ParserResult,
    ParserWarning,
)

logger = logging.getLogger(__name__)


class DocxAdapter(ParserAdapter):
    """Parser adapter for DOCX files using python-docx."""

    def name(self) -> str:
        return "python-docx"

    def version(self) -> str:
        try:
            import docx

            return getattr(docx, "__version__", "unknown")
        except ImportError:
            return "not installed"

    def capabilities(self) -> ParserCapabilities:
        return ParserCapabilities(
            supported_types=[FileType.DOCX],
            extracts_text=True,
            extracts_tables=True,
            extracts_layout=False,
            extracts_formulas=False,
            extracts_images=True,
            supports_ocr=False,
        )

    def parse(self, file_path: Path, file_type: FileType) -> ParserResult:
        result = ParserResult(
            parser_name=self.name(),
            parser_version=self.version(),
            file_type=file_type,
        )

        try:
            import docx
        except ImportError:
            result.errors.append("python-docx is not installed. Run: pip install python-docx")
            return result

        try:
            doc = docx.Document(str(file_path))
        except Exception as e:
            result.errors.append(f"Failed to open DOCX: {e}")
            return result

        text_blocks: list[dict[str, Any]] = []
        block_index = 0

        # Extract paragraphs
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            style_name = para.style.name if para.style else ""
            block_type = "paragraph"
            heading_level = 0
            font_size = 12.0

            if style_name.startswith("Heading"):
                block_type = "heading"
                try:
                    heading_level = int(style_name.split()[-1])
                except (ValueError, IndexError):
                    heading_level = 1
                # Approximate font size from heading level
                font_size = max(24 - (heading_level - 1) * 3, 12)
            elif style_name.startswith("List"):
                block_type = "list_item"
            elif style_name == "Title":
                block_type = "heading"
                heading_level = 1
                font_size = 28.0
            elif style_name == "Subtitle":
                block_type = "heading"
                heading_level = 2
                font_size = 20.0

            text_blocks.append({
                "text": text,
                "block_type": block_type,
                "heading_level": heading_level,
                "page_number": 1,
                "bbox": None,
                "font_size": font_size,
                "confidence": 1.0,
                "block_index": block_index,
            })
            block_index += 1

        # Build page data (DOCX has no page concept at API level — treat as single page)
        page_data: dict[str, Any] = {
            "page_number": 1,
            "width": 612,  # Letter size in points
            "height": 792,
            "text_blocks": text_blocks,
            "tables": [],
        }

        # Extract tables
        tables: list[dict[str, Any]] = []
        for table_idx, table in enumerate(doc.tables):
            cells: list[dict[str, Any]] = []
            row_count = len(table.rows)
            col_count = len(table.columns) if table.rows else 0

            for row_idx, row in enumerate(table.rows):
                for col_idx, cell in enumerate(row.cells):
                    cells.append({
                        "row_index": row_idx,
                        "column_index": col_idx,
                        "value": cell.text.strip(),
                        "data_type": "string",
                    })

            table_data: dict[str, Any] = {
                "page_number": 1,
                "title": f"Table {table_idx + 1}",
                "row_count": row_count,
                "column_count": col_count,
                "cells": cells,
                "confidence": 1.0,
            }
            tables.append(table_data)
            page_data["tables"].append(table_data)

        result.pages.append(page_data)
        result.text_blocks.extend(text_blocks)
        result.tables.extend(tables)
        result.page_count = 1

        # Extract images (presence only)
        figures: list[dict[str, Any]] = []
        try:
            from docx.opc.constants import RELATIONSHIP_TYPE as RT

            for rel in doc.part.rels.values():
                if "image" in rel.reltype:
                    figures.append({
                        "page_number": 1,
                        "caption": None,
                        "alt_text": None,
                        "bbox": None,
                    })
        except Exception:
            pass  # Image detection is best-effort

        result.figures.extend(figures)

        # Extract metadata from core properties
        try:
            props = doc.core_properties
            result.metadata = {
                "title": props.title or "",
                "author": props.author or "",
                "subject": props.subject or "",
                "created": str(props.created) if props.created else "",
                "modified": str(props.modified) if props.modified else "",
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables),
                "image_count": len(figures),
            }
        except Exception:
            result.metadata = {
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables),
                "image_count": len(figures),
            }

        if not text_blocks and not tables:
            result.warnings.append(
                ParserWarning(code="EMPTY_DOCUMENT", message="DOCX contains no text or tables.")
            )

        return result
